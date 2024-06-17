import pdfplumber
from langchain.text_splitter import TokenTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings
import os
from dotenv import load_dotenv
import docx
from io import BytesIO
from langchain.docstore.document import Document
import mammoth
import pandas as pd

load_dotenv()

embeddings = OpenAIEmbeddings()


def limpar_texto(texto):
    return " ".join(texto.split())


def extract_text_from_docx_with_mammoth(file):
    try:
        if isinstance(file, BytesIO):
            file.seek(0)
            result = mammoth.convert_to_html(file)
        else:
            with open(file, "rb") as docx_file:
                result = mammoth.convert_to_html(docx_file)
        return result.value
    except Exception as e:
        print(f"Erro ao extrair texto de arquivo DOCX: {e}")
        return ""


def process_document(files, file_extension):
    raw_text = ""
    try:
        if file_extension == "pdf":
            print("PROCESSANDO PDF")
            with pdfplumber.open(files) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        raw_text += text
                    for table in page.extract_tables():
                        for row in table:
                            raw_text += " ".join(str(row)) + "\n"
        elif file_extension in ["txt", "py", "cs", "java", "cpp", "js", "html", "css"]:
            print(f"PROCESSANDO {file_extension.upper()}")
            if isinstance(files, BytesIO):
                files.seek(0)
                raw_text += files.read().decode("utf-8")
        elif file_extension == "docx":
            print("PROCESSANDO DOCX")
            doc = docx.Document(files)
            raw_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        elif file_extension == "xlsx":
            print("PROCESSANDO XLSX")
            df = pd.read_excel(files)
            for col in df.columns:
                raw_text += " ".join(df[col].astype(str).values) + " "
        else:
            raise ValueError("Unsupported file type")

        text_splitter = TokenTextSplitter(chunk_size=350, chunk_overlap=80)
        documents = text_splitter.split_text(raw_text)
        document_objs = [Document(page_content=doc) for doc in documents]
        print(document_objs)
        return document_objs
    except Exception as e:
        print(f"Erro ao processar documento: {e}")
        return []


def criar_indices_faiss(files, file_extension):
    try:
        documents = process_document(files, file_extension)
        vector_store = FAISS.from_documents(documents, embeddings)
        vector_store.save_local("faiss_index")
        return vector_store
    except Exception as e:
        print(f"Erro ao criar índices FAISS: {e}")
        return None


def carregar_indices_faiss():
    try:
        vector_store = FAISS.load_local(
            "faiss_index", embeddings, allow_dangerous_deserialization=True
        )
        return vector_store
    except Exception as e:
        print(f"Erro ao carregar índices FAISS: {e}")
        return None


def adicionar_texto_ao_indice(vector_store, files, file_extension):
    try:
        documents = process_document(files, file_extension)
        # Ensure embeddings dimensionality matches the index
        document_embeddings = embeddings.embed_documents(
            [doc.page_content for doc in documents]
        )
        if document_embeddings and len(document_embeddings[0]) == vector_store.index.d:
            vector_store.add_documents(documents)
            vector_store.save_local("faiss_index")
        else:
            print(
                f"Dimension mismatch: embeddings dimension {len(document_embeddings[0])} vs index dimension {vector_store.index.d}"
            )
    except Exception as e:
        print(f"Erro ao adicionar texto ao índice: {e}")


def verificar_e_atualizar_indice(files, file_extension):
    try:
        if os.path.exists("faiss_index"):
            vector_store = carregar_indices_faiss()
            adicionar_texto_ao_indice(vector_store, files, file_extension)
        else:
            vector_store = criar_indices_faiss(files, file_extension)
        return vector_store
    except Exception as e:
        print(f"Erro ao verificar e atualizar índice: {e}")
        return None
